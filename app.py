import os
import streamlit as st
from io import BytesIO
from utils.file_parser import extract_text_from_docx, extract_text_from_pdf
from utils.generator import generate_all
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re

st.set_page_config(page_title="AI Requirements Analyzer", layout="wide")

# Initialize session state for authentication
if 'authenticated' not in st.session_state:
    st.session_state.authenticated = False

def check_password():
    """Check if user is authenticated"""
    # Get password from secrets or environment variable
    try:
        correct_password = st.secrets.get("APP_PASSWORD", os.environ.get("APP_PASSWORD", "admin123"))
    except:
        correct_password = os.environ.get("APP_PASSWORD", "admin123")
    
    # If already authenticated, return True
    if st.session_state.authenticated:
        return True
    
    # Show login form
    st.title("🔒 AI Requirements Analyzer - Login")
    st.markdown("Please enter the password to access the application.")
    
    col1, col2, col3 = st.columns([1, 2, 1])
    with col2:
        password = st.text_input("Password", type="password", label_visibility="collapsed", placeholder="Enter password")
        
        if st.button("Login", use_container_width=True):
            if password == correct_password:
                st.session_state.authenticated = True
                st.rerun()
            else:
                st.error("❌ Incorrect password. Please try again.")
    
    return False

def format_text_for_docx(text, doc):
    """Parse and format text with better styling for Word document"""
    if not text:
        return
    
    lines = text.split('\n')
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        
        # Skip empty lines (we'll add spacing manually)
        if not line:
            i += 1
            continue
        
        # Check for Epic (various formats: ## Epic, **Epic**, Epic:, etc.)
        epic_match = None
        epic_name = None
        
        patterns = [
            (r'^##?\s*(Epic|EPIC)\s*:?\s*(.+)', 2),
            (r'^\*\*(Epic|EPIC)\*\*\s*:?\s*(.+)', 2),
            (r'^(Epic|EPIC)\s*:?\s*(.+)', 2),
        ]
        
        for pattern, group_num in patterns:
            match = re.match(pattern, line, re.IGNORECASE)
            if match:
                epic_match = match
                epic_name = match.group(group_num).strip('*').strip()
                break
        
        if epic_match and epic_name:
            p = doc.add_paragraph()
            run = p.add_run(f'Epic: {epic_name}')
            run.bold = True
            run.font.size = Pt(12)
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
        
        # Check for User Story (various formats)
        elif any(re.match(p, line, re.IGNORECASE) for p in [
            r'^###?\s*(User Story|US|Story)\s*:?\s*(.+)',
            r'^\*\*(User Story|US|Story)\*\*\s*:?\s*(.+)',
            r'^(User Story|US|Story)\s*:?\s*(.+)',
            r'^(US-\d+[:\s]+)(.+)'
        ]):
            story_title = line
            # Clean up the title
            for pattern in [
                r'^###?\s*(User Story|US|Story)\s*:?\s*',
                r'^\*\*(User Story|US|Story)\*\*\s*:?\s*',
                r'^(User Story|US|Story)\s*:?\s*',
                r'^(US-\d+[:\s]+)'
            ]:
                story_title = re.sub(pattern, '', story_title, flags=re.IGNORECASE)
            story_title = story_title.strip('*').strip()
            
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.25)  # Indent user stories
            run = p.add_run(f'User Story: {story_title}')
            run.bold = True
            run.font.size = Pt(11)
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(6)
        
        # Check for bullet points or numbered lists
        elif re.match(r'^[-*•]\s+', line) or re.match(r'^\d+[.)]\s+', line):
            # Remove the bullet/number marker and add as formatted list
            clean_line = re.sub(r'^[-*•]\s+', '', line)
            clean_line = re.sub(r'^\d+[.)]\s+', '', clean_line)
            
            p = doc.add_paragraph(clean_line, style='List Bullet')
            p.paragraph_format.left_indent = Inches(0.5)
            p.paragraph_format.space_after = Pt(3)
        
        # Check for bold text in markdown (**text**)
        elif '**' in line:
            p = doc.add_paragraph()
            parts = re.split(r'(\*\*.*?\*\*)', line)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part.strip('*'))
                    run.bold = True
                else:
                    p.add_run(part)
            p.paragraph_format.space_after = Pt(6)
        
        # Regular paragraph
        else:
            p = doc.add_paragraph(line)
            p.paragraph_format.space_after = Pt(6)
        
        i += 1

# Check password before showing the app
if not check_password():
    st.stop()

# Main app content (only shown if authenticated)
# Header with title and logout button in top right
col1, col2 = st.columns([4, 1])
with col1:
    st.title("AI Requirements Analyzer — Streamlit")
with col2:
    st.markdown("<div style='height: 20px;'></div>", unsafe_allow_html=True)  # Vertical alignment
    if st.button("🔓 Logout", use_container_width=True, key="logout_btn", type="secondary"):
        st.session_state.authenticated = False
        st.rerun()

# st.markdown("Paste a business requirement or upload a doc (docx/pdf). The app will generate user stories, ACs, test cases and a mermaid flowchart using an LLM.")
st.markdown("Paste a business requirement or upload a doc (docx/pdf). The app will generate user stories, ACs, test cases and a summary using an LLM.")

# Read API key from Streamlit secrets or env
OPENAI_KEY = None
try:
    OPENAI_KEY = st.secrets["OPENAI_API_KEY"]
except Exception:
    OPENAI_KEY = os.environ.get("OPENAI_API_KEY")

if not OPENAI_KEY:
    st.warning("No OPENAI_API_KEY found. Set OPENAI_API_KEY as an env var or Streamlit secret before using the app.")

uploaded_file = st.file_uploader("Upload .docx or .pdf (optional)", type=["docx", "pdf", "txt"]) 
text_input = st.text_area("Or paste requirement text here:", height=200)

if uploaded_file is not None and text_input.strip() == "":
    # read bytes and detect
    bytes_data = uploaded_file.read()
    if uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        try:
            text_input = extract_text_from_docx(bytes_data)
        except Exception as e:
            st.error(f"Error reading docx: {e}")
    elif uploaded_file.type == "application/pdf":
        try:
            text_input = extract_text_from_pdf(bytes_data)
        except Exception as e:
            st.error(f"Error reading pdf: {e}")
    elif uploaded_file.type.startswith("text"):
        try:
            text_input = bytes_data.decode("utf-8")
        except UnicodeDecodeError:
            st.error("Error reading text file: Unable to decode as UTF-8. Please ensure the file is UTF-8 encoded.")
            text_input = ""

col1, col2 = st.columns([3,1])
with col2:
    model_choice = st.selectbox("Model", ["gpt-4o-mini"], index=0)
    max_tokens = st.slider("Max tokens (per section)", min_value=256, max_value=3000, value=1200, step=128)

if st.button("Generate Requirements"):
    if not text_input.strip():
        st.error("Please paste a requirement or upload a document.")
    elif not OPENAI_KEY:
        st.error("OpenAI API key not found. Set OPENAI_API_KEY in env or Streamlit secrets.")
    else:
        with st.spinner("Generating..."):
            try:
                outputs = generate_all(text_input, prompts_dir="prompts", api_key=OPENAI_KEY, max_tokens=max_tokens)
            except Exception as e:
                st.error(f"Error generating requirements: {e}")
                st.stop()

        st.success("Generation complete")

        st.header("Executive Summary")
        st.markdown(outputs.get('summary',''))

        st.header("Epics & User Stories")
        st.markdown(outputs.get('user_stories',''))

        st.header("Acceptance Criteria")
        st.markdown(outputs.get('acceptance_criteria',''))

        st.header("Test Cases")
        st.markdown(outputs.get('test_cases',''))

        # Mermaid Flowchart - Commented out
        # st.header("Mermaid Flowchart")
        # mermaid_text = outputs.get('flows','')
        # # render mermaid in an HTML component
        # mermaid_html = f"""
        # <script src=\"https://cdn.jsdelivr.net/npm/mermaid/dist/mermaid.min.js\"></script>
        # <div class=\"mermaid\">{mermaid_text}</div>
        # <script>mermaid.initialize({{startOnLoad:true}});</script>
        # """
        # st.components.v1.html(mermaid_html, height=300)

        # Create downloadable docx with improved formatting
        doc = Document()
        
        # Title
        title = doc.add_heading('AI Generated Requirements', level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Executive Summary
        doc.add_heading('Executive Summary', level=2)
        summary_text = outputs.get('summary', '')
        if summary_text:
            p = doc.add_paragraph(summary_text)
            p.paragraph_format.space_after = Pt(12)
        
        # Epics & User Stories with better formatting
        doc.add_heading('Epics & User Stories', level=2)
        user_stories_text = outputs.get('user_stories', '')
        if user_stories_text:
            format_text_for_docx(user_stories_text, doc)
        
        # Acceptance Criteria
        doc.add_heading('Acceptance Criteria', level=2)
        acceptance_text = outputs.get('acceptance_criteria', '')
        if acceptance_text:
            format_text_for_docx(acceptance_text, doc)
        
        # Test Cases
        doc.add_heading('Test Cases', level=2)
        test_cases_text = outputs.get('test_cases', '')
        if test_cases_text:
            format_text_for_docx(test_cases_text, doc)
        # Mermaid Flowchart - Commented out
        # doc.add_heading('Mermaid Flowchart', level=2)
        # doc.add_paragraph(outputs.get('flows',''))

        # Save to bytes buffer instead of file
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        fname = f"requirements_{int(datetime.now().timestamp())}.docx"
        st.download_button(
            label="Download as Word (.docx)",
            data=buffer.getvalue(),
            file_name=fname,
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )